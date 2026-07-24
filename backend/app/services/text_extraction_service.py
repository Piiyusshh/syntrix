from pathlib import Path
from typing import cast

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from fastapi import HTTPException

from app.schemas.extraction import (
    ExtractedDocument,
    ExtractedPage,
)


SUPPORTED_FILE_TYPES = {
    "pdf",
    "docx",
    "txt",
}


def extract_document(file_path: str) -> ExtractedDocument:
    """
    Extract text from a supported document while preserving page information.
    """

    path = Path(file_path)

    if not path.exists():
        raise HTTPException(
            status_code=404,
            detail="Document file not found.",
        )

    extension = path.suffix.lower().lstrip(".")

    if extension not in SUPPORTED_FILE_TYPES:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type: {extension}",
        )

    try:
        if extension == "pdf":
            pages = extract_pdf(path)

        elif extension == "docx":
            pages = extract_docx(path)

        else:
            pages = extract_txt(path)

        return ExtractedDocument(pages=pages)

    except HTTPException:
        raise

    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to extract text: {exc}",
        )


def extract_pdf(path: Path) -> list[ExtractedPage]:
    """
    Extract text from each page of a PDF.
    """

    pages: list[ExtractedPage] = []

    document = fitz.open(str(path))

    try:
        for page_index in range(document.page_count):
            page = document.load_page(page_index)

            page_text = cast(str, page.get_text("text"))

            pages.append(
                ExtractedPage(
                    page_number=page_index + 1,
                    text=clean_text(page_text),
                )
            )

        return pages

    finally:
        document.close()


def extract_docx(path: Path) -> list[ExtractedPage]:
    """
    Extract text from a DOCX document.
    """

    document = DocxDocument(str(path))

    text = "\n".join(
        paragraph.text
        for paragraph in document.paragraphs
    )

    return [
        ExtractedPage(
            page_number=1,
            text=clean_text(text),
        )
    ]


def extract_txt(path: Path) -> list[ExtractedPage]:
    """
    Extract text from a TXT document.
    """

    text = path.read_text(
        encoding="utf-8",
        errors="ignore",
    )

    return [
        ExtractedPage(
            page_number=1,
            text=clean_text(text),
        )
    ]


def clean_text(text: str) -> str:
    """
    Normalize extracted text by removing extra whitespace
    and empty lines.
    """

    lines = []

    for line in text.splitlines():
        cleaned = " ".join(line.split())

        if cleaned:
            lines.append(cleaned)

    return "\n".join(lines).strip()